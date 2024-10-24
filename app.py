from flask import Flask, request, jsonify, render_template, send_file
import fitz  # PyMuPDF for PDF processing
from transformers import pipeline
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document
import os

app = Flask(__name__)

# Load a pre-trained model for summarization
summarizer = pipeline("summarization", model="sshleifer/distilbart-cnn-12-6")

# In-memory store for bookmarks
bookmarks = {}

# Function to extract text from PDF
def extract_text_from_pdf(pdf_path):
    pdf = fitz.open(pdf_path)
    text = ""
    for page in pdf:
        text += page.get_text()
    return text

# Function to handle large text by splitting it into chunks
def summarize_large_text(text, max_chunk_size=1000):
    # Split the text into smaller chunks
    text_chunks = [text[i:i + max_chunk_size] for i in range(0, len(text), max_chunk_size)]
    
    summaries = []
    for chunk in text_chunks:
        summary = summarizer(chunk, max_length=150, min_length=50, do_sample=False)
        summaries.append(summary[0]['summary_text'])
    
    # Combine all summaries
    return " ".join(summaries)

# Route to upload and summarize PDF
@app.route('/upload', methods=['POST'])
def upload_pdf():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400

    # Save the file temporarily
    file_path = "./uploads/" + file.filename
    file.save(file_path)

    # Extract text from the PDF
    extracted_text = extract_text_from_pdf(file_path)

    if not extracted_text:
        return jsonify({"error": "No text extracted from the PDF"}), 400

    # Summarize the extracted text
    summary = summarize_large_text(extracted_text)

    return jsonify({"summary": summary, "extracted_text": extracted_text})

# Route to add a bookmark
@app.route('/add_bookmark', methods=['POST'])
def add_bookmark():
    data = request.get_json()
    user_id = data['user_id']
    section = data['section']
    
    if user_id not in bookmarks:
        bookmarks[user_id] = []
    
    bookmarks[user_id].append(section)
    
    return jsonify({"message": "Bookmark added successfully", "bookmarks": bookmarks[user_id]})

# Route to retrieve bookmarks for a user
@app.route('/get_bookmarks', methods=['POST'])
def get_bookmarks():
    data = request.get_json()
    user_id = data['user_id']
    
    if user_id in bookmarks:
        return jsonify({"bookmarks": bookmarks[user_id]})
    else:
        return jsonify({"bookmarks": []})

# Function to export content as a PDF
def export_to_pdf(content, filename):
    pdf_path = f"./exports/{filename}.pdf"
    c = canvas.Canvas(pdf_path, pagesize=letter)
    width, height = letter
    c.drawString(100, height - 100, content)
    c.save()
    return pdf_path
# Function to export content as a Word document
def export_to_word(content, filename):
    # Ensure the filename ends with .docx
    if not filename.endswith('.docx'):
        filename += '.docx'

    doc = Document()
    doc.add_heading('Extracted Content', 0)
    doc.add_paragraph(content)
    word_path = f"./exports/{filename}"
    doc.save(word_path)
    return word_path

# Route to export extracted content
@app.route('/export', methods=['POST'])
def export_content():
    data = request.get_json()
    content = data['content']
    format_type = data['format']
    filename = data['filename']

    if format_type == 'pdf':
        pdf_path = export_to_pdf(content, filename)
        return send_file(pdf_path, as_attachment=True)
    elif format_type == 'word':
        word_path = export_to_word(content, filename)
        return send_file(word_path, as_attachment=True)
    else:
        return jsonify({"error": "Unsupported format"}), 400


# Basic route for the homepage
@app.route('/')
def index():
    return render_template('index.html')

if __name__ == '__main__':
    # Create uploads and exports folders if they don't exist
    if not os.path.exists("./uploads"):
        os.makedirs("./uploads")
    if not os.path.exists("./exports"):
        os.makedirs("./exports")
    app.run(debug=True)
